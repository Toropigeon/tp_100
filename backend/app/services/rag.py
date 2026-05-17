from pathlib import Path
import re
import zipfile


class DocumentationIndex:
    def __init__(self, docs_dir: Path):
        self.docs_dir = docs_dir
        self.chunks = self._load_chunks()

    def search(self, query: str, limit: int = 5) -> list[str]:
        query_terms = set(_tokenize(query))
        if not query_terms:
            return self.chunks[:limit]

        scored: list[tuple[int, str]] = []
        for chunk in self.chunks:
            score = len(query_terms.intersection(_tokenize(chunk)))
            if score:
                scored.append((score, chunk))

        scored.sort(key=lambda item: item[0], reverse=True)
        return [chunk for _, chunk in scored[:limit]] or self.chunks[:limit]

    def _load_chunks(self) -> list[str]:
        texts: list[str] = []
        for path in sorted(self.docs_dir.glob("*")):
            if path.suffix.lower() == ".docx":
                texts.append(_read_docx_text(path))
            elif path.suffix.lower() == ".pdf":
                texts.append(_read_pdf_text(path))

        joined = "\n".join(text for text in texts if text.strip())
        chunks = _chunk_text(joined, size=900)
        return chunks or ["Документация не была прочитана. Используй инженерную интерпретацию по диагностическим данным."]


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document

        document = Document(path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(paragraphs)
    except Exception:
        try:
            with zipfile.ZipFile(path) as archive:
                xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            return re.sub(r"<[^>]+>", " ", xml)
        except Exception:
            return ""


def _read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""


def _chunk_text(text: str, size: int) -> list[str]:
    clean = re.sub(r"\s+", " ", text).strip()
    if not clean:
        return []
    return [clean[index : index + size] for index in range(0, len(clean), size)]


def _tokenize(text: str) -> list[str]:
    return re.findall(r"[a-zа-я0-9]{3,}", text.lower().replace("ё", "е"))

